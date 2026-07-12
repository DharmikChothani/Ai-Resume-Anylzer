import fitz
from docx import Document

def extract_pdf(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"Error reading PDF: {e}"

def extract_docx(file):
    try:
        doc = Document(file)
        text = ""
        
        # 1. Extract standard paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # 2. Extract text hidden inside tables (Crucial for resumes!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
    except Exception as e:
        return f"Error reading Word Document: {e}"

def extract_text(uploaded_file):
    # Make sure we handle uppercase extensions like .PDF or .DOCX
    filename = uploaded_file.name.lower()
    
    if filename.endswith(".pdf"):
        return extract_pdf(uploaded_file)
    elif filename.endswith(".docx"):
        return extract_docx(uploaded_file)
    else:
        # Return a clean error instead of crashing
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX.")